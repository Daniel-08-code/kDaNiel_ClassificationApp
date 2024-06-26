from flask import Flask, render_template, request
import matplotlib.pyplot as plt
import base64
from docx.shared import Inches
from PIL import Image
import io
from scipy.cluster import hierarchy
import numpy as np
import pandas as pd
import matplotlib
from docx import Document
import os
import math
import random
matplotlib.use('Agg') 
app = Flask(__name__)

save_path = os.path.join(app.root_path, 'static')


class Centroid:
    def __init__(self, Individu):
        self.data = Individu.data
        
class Individu:
    def __init__(self,data):
        self.group = 0
        self.data = data # data est les données de la ligne en question de dataTable
        
class Cluster:
    nbrCluster = 0

    def __init__(self, data):
        Cluster.nbrCluster += 1
        self.numeroCluster = Cluster.nbrCluster
        self.indicesIndividus = [Cluster.nbrCluster]  # Initialise avec un tableau contenant seulement le numéro de cluster
        self.data = data


class Fonction:
    @staticmethod
    def distanceEuclidienne(arg1, arg2):
        somme = 0.0
        for i in range(len(arg1.data)):
            somme += (arg1.data[i] - arg2.data[i]) ** 2
        return np.sqrt(somme)
    
    @staticmethod
    def extract_data_from_excel(file_path):
        # Charger le fichier Excel
        dataframe = pd.read_excel(file_path, engine='openpyxl')
        # Identifier la position du tableau de données en recherchant les index des premières cellules non vides
        start_row = None
        start_col = None
        for i in range(dataframe.shape[0]):
            for j in range(dataframe.shape[1]):
                if pd.notna(dataframe.iloc[i, j]):
                    start_row = i
                    start_col = j
                    break
            if start_row is not None:
                break
        # Collecter les noms des individus (ligne d'index 0)
        individus = dataframe.iloc[start_row+1:, start_col].tolist()
        # Collecter les noms des variables (première ligne après les noms des individus)
        variables = dataframe.iloc[start_row, start_col+1:].tolist()
        # Extraire les valeurs dans un tableau à deux dimensions
        values = dataframe.iloc[start_row+1:, start_col+1:]
        return individus, variables, values


    @staticmethod
    def distanceManatthan(arg1, arg2):
        somme = 0.0
        for i in range(len(arg1.data)):
            somme += abs(arg1.data[i] - arg2.data[i])
        return np.sqrt(somme)
    
    @staticmethod
    def aleatoire(k, tabIndividus):
        centroidTab = []
        indiceTab = []
        for i in range(0,k):
            indice = random.randint(0,len(tabIndividus)-1)
            while indice in indiceTab:
                indice = random.randint(0,len(tabIndividus)-1)
            indiceTab.append(indice)
            centroidTab.append(Centroid(tabIndividus[indice]))
        return centroidTab,indiceTab
    
    @staticmethod
    def mean2DTab(tab):
        abscisse, ordonnee = [0,0]
        if len(tab) % 2 == 0:
            for i in range(0,len(tab)):
                if i < len(tab)/2:
                    abscisse += tab[i]/len(tab)/2
                else:
                    ordonnee += tab[i]/len(tab)/2
        else:
            for i in range(0,len(tab)):
                if i < (len(tab)+1)/2:
                    abscisse += tab[i]/(len(tab)-1)/2
                else:
                    ordonnee += tab[i]/(len(tab)+1)/2
        return abscisse, ordonnee
    
    @staticmethod
    def kmeans(k,tabIndividus,distanceChoisie):
        doc = Document(os.path.join(save_path, 'rapport.docx'))
        doc.add_page_break()
        doc.add_paragraph("\n\n\n")
        centroidTab,indiceTab = Fonction.aleatoire(k, tabIndividus) 
        print(indiceTab)
        #choix aléatoire des centoides initiaux
        nbreModification = 10
        while nbreModification != 0 :
            nbreModification = 0
            #Affectation des individus aux groupes
            for i in range(0,len(tabIndividus)):
                if distanceChoisie == 'distanceEuclidienne':
                    distMin = Fonction.distanceEuclidienne(tabIndividus[i],centroidTab[tabIndividus[i].group]) #distance entre l'individu et son ancien centre
                    for j in range(0,len(centroidTab)):
                        if distMin > Fonction.distanceEuclidienne(tabIndividus[i],centroidTab[j]):
                            distMin = Fonction.distanceEuclidienne(tabIndividus[i],centroidTab[j])
                            tabIndividus[i].group = j
                            nbreModification += 1
                elif  distanceChoisie == 'distanceManatthan':
                    distMin = Fonction.distanceManatthan(tabIndividus[i],centroidTab[tabIndividus[i].group]) #distance entre l'individu et son ancien centre
                    for j in range(0,len(centroidTab)):
                        if distMin > Fonction.distanceManatthan(tabIndividus[i],centroidTab[j]):
                            distMin = Fonction.distanceManatthan(tabIndividus[i],centroidTab[j])
                            tabIndividus[i].group = j
                            nbreModification += 1
            if nbreModification != 0 : # cas ou les groupes ont changé
                #Calcul des nouveaux centres
                for i in range(k):  # k est le nombre de centres
                    centroid_data_sum = np.zeros_like(centroidTab[i].data)  # Initialiser la somme des données des individus pour chaque centre
                    ctr = np.zeros(len(centroidTab[i].data))  # Initialiser le compteur pour chaque dimension
                    
                    # Calculer la somme des données des individus pour chaque centre
                    for l in range(len(tabIndividus)):
                        if tabIndividus[l].group == i:
                            centroid_data_sum += tabIndividus[l].data
                            ctr += 1
                    
                    # Mettre à jour les centres en calculant la moyenne des données des individus pour chaque centre
                    centroidTab[i].data = np.where(ctr != 0, centroid_data_sum / ctr, centroidTab[i].data)

        TabCenters = [centroidTab[i].data for i in range(len(centroidTab))]
        resultat = []   #resultat sous forme de liste indiquant le groupe auquel chaque individu appartient
        for i in range(0,len(tabIndividus)):
            resultat.append(tabIndividus[i].group)
            doc.add_paragraph(f"{i+1}-- Individu {i+1} appartient au groupe {tabIndividus[i].group}.\n")
        doc.save(os.path.join(save_path, 'telechargeable.docx'))
        return resultat,TabCenters

    @staticmethod
    def hierarchy(clusters,distanceChoisie):
        doc = Document(os.path.join(save_path, 'rapport.docx'))
        doc.add_page_break()
        doc.add_paragraph("\n\n\n")
        hierarchy = []
        indices = []
        iteration = 0
        while len(clusters)-2*iteration > 1:
            distanceMin = float('inf')
            indice1, indice2 = None, None
            for i in range(len(clusters)):
                if i not in indices:
                    for j in range(i + 1, len(clusters)):
                        if j not in indices:
                            if distanceChoisie == 'distanceEuclidienne':
                                distance = Fonction.distanceEuclidienne(clusters[i], clusters[j])
                            elif  distanceChoisie == 'distanceManatthan':
                                distance = Fonction.distanceManatthan(clusters[i], clusters[j])
                            if distance < distanceMin:
                                distanceMin = distance
                                indice1, indice2 = i, j
            indices.append(indice1)
            indices.append(indice2)
            iteration +=1
            new_data = [(clusters[indice1].data[i] + clusters[indice2].data[i]) / 2 for i in range(len(clusters[indice1].data))]
            newCluster = Cluster(new_data)  # Crée un nouveau cluster avec les données du barycentre des deux clusters
            newCluster.indicesIndividus = clusters[indice1].indicesIndividus + clusters[indice2].indicesIndividus
            clusters.append(newCluster)
            doc.add_paragraph(f"{iteration}-- Liaison du Cluster {clusters[indice1].numeroCluster-1} et du Cluster {clusters[indice2].numeroCluster-1} pour former le cluster {newCluster.numeroCluster-1} regroupant ainsi {len(newCluster.indicesIndividus)} elements.\n")
    # Enregistrer le document Word dans le dossier static
            doc.save(os.path.join(save_path, 'telechargeable.docx'))
            hierarchy.append([indice1, indice2, distanceMin, len(newCluster.indicesIndividus)])
        return hierarchy
    
    
@app.route('/', methods=['GET', 'POST'])
def upload_and_display():
    if request.method == 'POST':
        method = request.form['method']
        uploaded_file = request.files['file']
        distance = request.form['distance']
        
        if uploaded_file.filename != '':
            individus, variables, data = Fonction.extract_data_from_excel(uploaded_file)
            columns = data.columns.tolist()
            rows = data.values.tolist()
            
            data = data.to_numpy()
            data2D = [Fonction.mean2DTab(data[i]) for i in range(len(data))]
            data2D = np.array(data2D)
            if method == 'CAH':
                Cluster.nbrCluster = 0
                clusters = [Cluster(data[i]) for i in range(len(data))]
                
        
                hierarchy_result = Fonction.hierarchy(clusters,distance)
                hierarchy_result = np.array(hierarchy_result)
                
                #On peut verifier le resultat avec ce package
                #Z = hierarchy.linkage(data, method='centroid') 
                #print(hierarchy_result)
            
                # Création du dendrogramme
                plt.figure(figsize=(10, 8))
                hierarchy.dendrogram(hierarchy_result, labels=individus)
                plt.xticks(rotation=90)
                plt.title('CAH Clustering by Daniel ANONWODJI',fontsize=15)
                plt.xlabel('Noms des individus')
                plt.ylabel('Distance')

                # Sauvegarde de la figure dans un tampon mémoire
                buffer = io.BytesIO()
                plt.savefig(buffer, format='png')
                buffer.seek(0)

                # Conversion de la figure en format base64
                image_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                # Fermeture de la figure pour libérer la mémoire
                plt.close()
                
            elif method == 'K-Means':
                tabIndividus = [Individu(data[i]) for i in range(len(data))]
                num_clusters = int(request.form['kmeans-groups'])
                resultat,centers = Fonction.kmeans(num_clusters, tabIndividus,distance)
                centers = np.array(centers)
                print(centers)
                centers2D = [Fonction.mean2DTab(centers[i]) for i in range(len(centers))]
                centers2D = np.array(centers2D)
                resultat = np.array(resultat)
                print(resultat)
                # Plotting the clusters (you can customize this part according to your visualization needs)
                plt.figure(figsize=(20, 20))
                plt.scatter(data2D[:, 0], data2D[:, 1], c=resultat, cmap='viridis', s=1500, alpha=1)
                plt.title('K-Means Clustering by Daniel ANONWODJI',fontsize=30)
                plt.xlabel('Feature 1')
                plt.ylabel('Feature 2')
                plt.scatter(centers2D[:, 0], centers2D[:, 1], c='red', s=300, alpha=1, marker='x', label='Centroids')
                for i, (x, y) in enumerate(zip(data2D[:, 0], data2D[:, 1])):
                    plt.text(x, y, individus[i], fontsize=25, ha='right', va='bottom')
                plt.legend(fontsize=20,loc='lower left')
                # Saving the plot as a base64 encoded image
                buffer = io.BytesIO()
                plt.savefig(buffer, format='png')
                buffer.seek(0)
                image_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                plt.close()
            image_data = base64.b64decode(image_base64)
            image = Image.open(io.BytesIO(image_data))
            doc = Document(os.path.join(save_path, 'telechargeable.docx'))
            doc.add_picture(io.BytesIO(image_data), width=Inches(6))
            doc.save(os.path.join(save_path, 'telechargeable.docx'))
            return render_template('index.html', columns=columns, rows=rows, image=image_base64)
        else:
            return render_template('index.html')
    else:
        return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)



